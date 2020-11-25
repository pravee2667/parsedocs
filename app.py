# -*- coding: utf-8 -*-

import spacy
#import re
#import sys
from spacy.matcher import Matcher 

from docx.api import Document
from bin import extractor as ex
import logging
from bin import linkedin
from bin import pdf_extractor as pd
from bin import loader
from flask import Flask,request,session,render_template
import os
from outlook import meeting
from werkzeug.utils import secure_filename
import json


app=Flask(__name__)
app.secret_key="sfsdfdsfdsf"
nlp = spacy.load('en_core_web_sm')

matcher = Matcher(nlp.vocab)   

#Home Page 
@app.route('/')
def home():
    return render_template('home.html')

#Upload files
@app.route('/upload',methods=['POST','GET'])
def parse_doc():
    if request.method=='POST':
        if not request.form:
            filename=request.files['myfile']
            upload_dir=os.path.join(app.root_path,'Data','PDF')
            print(upload_dir)
            txt=list()
            top_skills=None
            logging.getLogger().setLevel(logging.INFO)
            logging.info('File Name {}'.format(filename))
            fileextens=filename.filename.split('.')[1]
            logging.info('File Extension {}'.format(fileextens))
            if filename.filename.endswith('docx'):
                logging.info('File name ends with docx')
                doc=Document(filename)
                for para in doc.paragraphs:
                    txt.append(para.text)
                full_t= ' '.join(txt)
                
                name_e=ex.name_extraction(full_t)
                name_dup=name_e 
                logging.info('Extracted Name {}'.format(name_dup))
                
                mob=ex.extract_mob_number(full_t)
                mob_dup=mob
                logging.info('Extracted Mobile  {}'.format(mob_dup))
                
                mail=ex.extract_mail(full_t)
                mail_dup=mail 
                logging.info('Extracted mail {}'.format(mail))
                
                skills=ex.extract_skills(doc)
                logging.info('Extracted Skills {}'.format(skills))
                #skills=[]
                
                try:
                    top_skills=",".join(skills[0:3])
                    skills_dup=skills[0:3]
                except Exception as exx:
                    top_skills=None
                
                if not mail or not name_e  or not mob  or not skills: 
                    txt_doc=ex.txt_extraction(filename)
                    
                if not mob:
                    mob_doc=ex.extract_mob_number(txt_doc)
                    
                if not mail:
                    mail=ex.extract_mail(txt_doc)
                    
                if not name_e:
                    name_doc=ex.name_extraction(txt_doc)
                    name_e=name_doc
                if not skills:
                    logging.info('Extracting Skills from text file {}'.format(name_e))
                    url=ex.extract_linkedinurl(txt_doc)
                    #url='https://www.linkedin.com/in/praveensk-kumarsk227'
                    if url:
                        skills=linkedin.skills_linkdn(url)
                        if skills:
                            top_skills=",".join(skills[0:3])
                            skills_dup=skills[0:3]
                    else:
                        skill_set=pd.pdf_extract_skills(txt_doc)
                        skills_list=[]
                        try:
                            for i in skill_set:
                                skills_list.append(i)
                            top_skills=",".join(skills_list)
                        except Exception as exx:
                            top_skills=None
                        print("top {}".format(top_skills))
                        skills_dup=skills_list
#                session['name_session']=name_dup
#                session['mob_session']=mob_dup
                session['mail_session']=mail_dup
                session['skills_session']=skills_dup
                logging.info('Extracted Skills {}'.format(skills_dup))
                ret=list()
                ret.append([name_e,mail_dup,skills_dup])
                
                return str(ret)
                #return render_template('home.html',output=mail,Username=name_e,skills_list=top_skills)
            elif fileextens=='pdf':
                secure_file=secure_filename(filename.filename)
                filename.save(os.path.join(upload_dir,secure_file))
                files=upload_dir+'/'+secure_file
                logging.info('PDF file {}'.format(files))
                
                pdftext=pd.pdf_txt(files)
                pdftx="".join(pdftext)
                name_e=ex.name_extraction(pdftx)
                logging.info('Extracted Name {}'.format(name_e))
                
                mob=ex.extract_mob_number(pdftx)
                #session['mob_session']=mob
                logging.info('Extracted Mobile  {}'.format(mob))
                
                mail=ex.extract_mail(pdftx)
                logging.info('Extracted mail {}'.format(mail))
                mail_dup=mail
                
                skills=pd.pdf_extract_skills(pdftext)
                logging.info('Extracted Skills {}'.format(skills))
                
                skills_list=[]
                try:
                    for i in skills:
                        skills_list.append(i)
                    top_skills=",".join(skills_list)
                except Exception as exx:
                    top_skills=None
                print("top {}".format(top_skills))
                if not skills:
                    
                    url=ex.extract_linkedinurl(pdftx)
                    if url:
                        skills=linkedin.skills_linkdn(url)
                        skills_list=skills
                        if skills:
                            top_skills=",".join(skills[0:3])
                    else:
                        top_skills=None
                session['mail_session']=mail_dup
                session['skills_session']=skills_list
                ret=list()
                ret.append([name_e,mail_dup,skills_dup])
                print(ret)
                return str(ret)
                #return render_template('home.html',output=mail,Username=name_e,skills_list=top_skills)          
        else:
            if request.form['submitbutton']:
                ski=session.get('skills_session')
                logging.info('Extracted Skills {}'.format(ski))
                if ski:
                    roles=meeting.Get_roles(ski)
                    actual_roles=json.loads(roles.text)
                    logging.info('Extracted Roles {}'.format(actual_roles))
                    if actual_roles:
                        role1=actual_roles[0]
                        print(role1)
                        logging.info('Extracted Roless {}'.format(actual_roles))
                    else:
                        actual_roles=['DataScientist']
                    
                    return str(actual_roles)
                    #return render_template('home.html',Roles=actual_roles)
                else:
                    ss=request.form['Skills']
                
                    skills_listt=ss.split(',')
                    logging.info('Extracted Input Skills {}'.format(skills_listt))
                    roles=meeting.Get_roles(skills_listt)
                    logging.info('Extracted Roles {}'.format(roles.text))
                    
                    actual_roles=json.loads(roles.text)
                    role1=actual_roles[0]
                    print(role1)
                    logging.info('Extracted Roless {}'.format(actual_roles))
                    return str(actual_roles)
                    #return render_template('home.html',Roles=actual_roles)
                    
            elif request.form['Roles']:
                return "Success"
            else:
                return "Failed"
            
     
    return render_template('home.html')

@app.route('/roless',methods=['POST','GET'])
def free_slots():
    if request.method=='POST':
        role=request.form['rol']
        print("Roles")
        print(role)
#        role_list=json.loads(role)
#        print(role_list)
        time_slot=meeting.Get_slots(role)
        session['role']=role
        time_slot_list=json.loads(time_slot)
        slot_top=time_slot_list[0:3]
        print(str(slot_top))
        return str(slot_top)
        #return  render_template('home.html',data=slot_top)
    
@app.route('/slots',methods=['POST','GET'])    
def meeting_url():
    if request.method=='POST':
        slot=request.form['slots']
        print(slot)
        role=session['role']
        date=slot.split(',')[0]
        time=slot.split(',')[1]
        meeting_req=meeting.Schedule_meet(role,date,time)
        print(meeting_req)
        meeting.Send_email(meeting_req,'praveen.sabhiniveeshu@pacteraedge.com')
        return meeting_req
        
    
    


@app.route('/mail')            
def mail_extract():
    mail=session.get('mail_session')
    print(mail)
    return "mail"



    
#def execute(inp,txt=[]):
#    logging.getLogger().setLevel(logging.INFO)
#    filename=inp[1]
#    logging.info('File Name {}'.format(filename))
#    if filename.endswith('.docx'):
#        logging.info('File name ends with docx')
#        doc=Document(filename)
#        for para in doc.paragraphs:
#            txt.append(para.text)
#        full_t= ' '.join(txt)
#        name_e=ex.name_extraction(full_t)
#        logging.info('Extracted Name {}'.format(name_e))
#        
#        mob=ex.extract_mob_number(full_t)
#        logging.info('Extracted Mobile  {}'.format(mob))
#        
#        mail=ex.extract_mail(full_t)
#        logging.info('Extracted mail {}'.format(mail))
#        
#        skills=ex.extract_skills(doc)
#        logging.info('Extracted Skills {}'.format(skills))
#        
#        
#        if not mail or not name_e  or not mob  or not skills:   
#            txt_doc=ex.txt_extraction(filename)
#        if not mob:
#            mob_doc=ex.extract_mob_number(txt_doc)
#            print(mob_doc)
#        if mail=='Not Found':
#            mail_doc=ex.extract_mail(txt_doc)
#            print(mail_doc)
#        if not name_e:
#            name_doc=ex.name_extraction(txt_doc)
#            print(name_doc)
#        if not skills:
#            url=ex.extract_linkedinurl(txt_doc)
#            skills=linkedin.skills_linkdn(url)
#            print(skills)
            
            
            
            
      

        
#def txt_extraction(file):
#    temp = docx2txt.process(file)
#    text = [line.replace('\t', ' ') for line in temp.split('\n') if line]
#    full_t= ' '.join(text)
#    return full_t
#    
#    
#def name_extraction(resume_text):
#    nlp_text = nlp(resume_text)
#    
#    # First name and Last name are always Proper Nouns
#    pattern = [{'POS': 'PROPN'}, {'POS': 'PROPN'}]
#    
#    matcher.add('NAME', None, pattern)
#    
#    matches = matcher(nlp_text)
#
#    for match_id, start, end in matches:
#        span = nlp_text[start:end]
#        span.text
#    return nlp_text[0:2]
#
#
#
#def extract_mob_number(text):
#    phone = re.findall(re.compile(r'(?:(?:\+?([1-9]|[0-9][0-9]|[0-9][0-9][0-9])\s*(?:[.-]\s*)?)?(?:\(\s*([2-9]1[02-9]|[2-9][02-8]1|[2-9][02-8][02-9])\s*\)|([0-9][1-9]|[0-9]1[02-9]|[2-9][02-8]1|[2-9][02-8][02-9]))\s*(?:[.-]\s*)?)?([2-9]1[02-9]|[2-9][02-9]1|[2-9][02-9]{2})\s*(?:[.-]\s*)?([0-9]{4})(?:\s*(?:#|x\.?|ext\.?|extension)\s*(\d+))?'), text)
#    
#    if phone:
#        number = ''.join(phone[0])
#        if len(number) > 10:
#            return '+' + number
#        else:
#            return number
#
#
#def extract_mail(email):
#    email = re.findall("([^@|\s]+@[^@]+\.[^@|\s]+)", email)
#    #print(email)
#    if email:
#        try:
#            return email[0].split()[0].strip(';')
#        except IndexError:
#            return None
#
#def iter_headings(paragraphs):
#    for i in range(0,len(paragraphs)):
#        if paragraphs[i].style.name.startswith('Heading'):
#            yield paragraphs[i],i
#
#
#def extract_skills(doch,aft=[],skills=[]):
#    for heading,i in iter_headings(doch.paragraphs):
#        if 'Skills' in heading.text :
#            for j in range(i+1,len(doch.paragraphs)):
#                #print(doch.paragraphs[j].text)
#                aft.append(doch.paragraphs[j].text)
#                txt=doch.paragraphs[j].text
#                if doch.paragraphs[j].style.name.startswith('Heading'):
#                    break
#                if not txt:
#                    continue
#                else:
#                    if ':' in txt :
#                        word=txt.split(':')
#                        skills.append(word[1])
#                    else:
#                        skills.append(txt)    
#            return skills
#    

    
    
if __name__=="__main__":
    app.run(host='127.0.0.1', port=8080, debug=True)
    #execute(sys.argv)