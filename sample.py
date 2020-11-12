# -*- coding: utf-8 -*-
# -*- coding: utf-8 -*-

import spacy
#import re
#import sys
from spacy.matcher import Matcher 
from docx.api import Document
from bin import extractor as ex
import logging
from bin import linkedin

from flask import Flask,request,session,jsonify

app=Flask(__name__)
app.secret_key="sfsdfdsfdsf"
nlp = spacy.load('en_core_web_sm')

matcher = Matcher(nlp.vocab)   
 
@app.route('/home')
def home():
    return "PrepTalk Hiring"

@app.route('/upload',methods=['POST','GET'])
def parse_doc():
    if request.method=='POST':
        filename=request.files['img']
        #print(filename)
        txt=list()
        logging.getLogger().setLevel(logging.INFO)
        logging.info('File Name {}'.format(filename))
        fileextens=filename.filename.split('.')[1]
        print(fileextens)
        if fileextens=='docx':
            logging.info('File name ends with docx')
            doc=Document(filename)
            for para in doc.paragraphs:
                txt.append(para.text)
            full_t= ' '.join(txt)
            name_e=ex.name_extraction(full_t)
            session['name_session']=name_e
            logging.info('Extracted Name {}'.format(name_e))
            
            mob=ex.extract_mob_number(full_t)
            session['mob_session']=mob
            logging.info('Extracted Mobile  {}'.format(mob))
            
            mail=ex.extract_mail(full_t)
            session['mail_session']=mail
            logging.info('Extracted mail {}'.format(mail))
            print(session.get('mail_session'))
            skills=ex.extract_skills(doc)
            logging.info('Extracted Skills {}'.format(skills))
            session['skills_session']=skills
            
            
            if not mail or not name_e  or not mob  or not skills: 
                
                txt_doc=ex.txt_extraction(filename)
            if not mob:
                mob_doc=ex.extract_mob_number(txt_doc)
                print(mob_doc)
            if not mail:
                mail_doc=ex.extract_mail(txt_doc)
                print(mail_doc)
            if not name_e:
                name_doc=ex.name_extraction(txt_doc)
                print(name_doc)
            if not skills:
                
                url=ex.extract_linkedinurl(txt_doc)
                skills=linkedin.skills_linkdn(url)
                print(skills)
        print("skills")    
        return "Skills"
     
    return "File Not Uploaded"

@app.route('/mail')            
def mail_extract():
    mail=session.get('mail_session')
    print(mail)
    return "mail"
    
#def execute(inp,txt=[]):
#    logging.getLogger().setLevel(logging.INFO)
#    filename=inp[1]
#    logging.info('File Name {}'.format(filename))
#    if filename.endswith('.docx'):
#        logging.info('File name ends with docx')
#        doc=Document(filename)
#        for para in doc.paragraphs:
#            txt.append(para.text)
#        full_t= ' '.join(txt)
#        name_e=ex.name_extraction(full_t)
#        logging.info('Extracted Name {}'.format(name_e))
#        
#        mob=ex.extract_mob_number(full_t)
#        logging.info('Extracted Mobile  {}'.format(mob))
#        
#        mail=ex.extract_mail(full_t)
#        logging.info('Extracted mail {}'.format(mail))
#        
#        skills=ex.extract_skills(doc)
#        logging.info('Extracted Skills {}'.format(skills))
#        
#        
#        if not mail or not name_e  or not mob  or not skills:   
#            txt_doc=ex.txt_extraction(filename)
#        if not mob:
#            mob_doc=ex.extract_mob_number(txt_doc)
#            print(mob_doc)
#        if mail=='Not Found':
#            mail_doc=ex.extract_mail(txt_doc)
#            print(mail_doc)
#        if not name_e:
#            name_doc=ex.name_extraction(txt_doc)
#            print(name_doc)
#        if not skills:
#            url=ex.extract_linkedinurl(txt_doc)
#            skills=linkedin.skills_linkdn(url)
#            print(skills)
            
            
            
            
      

        
#def txt_extraction(file):
#    temp = docx2txt.process(file)
#    text = [line.replace('\t', ' ') for line in temp.split('\n') if line]
#    full_t= ' '.join(text)
#    return full_t
#    
#    
#def name_extraction(resume_text):
#    nlp_text = nlp(resume_text)
#    
#    # First name and Last name are always Proper Nouns
#    pattern = [{'POS': 'PROPN'}, {'POS': 'PROPN'}]
#    
#    matcher.add('NAME', None, pattern)
#    
#    matches = matcher(nlp_text)
#
#    for match_id, start, end in matches:
#        span = nlp_text[start:end]
#        span.text
#    return nlp_text[0:2]
#
#
#
#def extract_mob_number(text):
#    phone = re.findall(re.compile(r'(?:(?:\+?([1-9]|[0-9][0-9]|[0-9][0-9][0-9])\s*(?:[.-]\s*)?)?(?:\(\s*([2-9]1[02-9]|[2-9][02-8]1|[2-9][02-8][02-9])\s*\)|([0-9][1-9]|[0-9]1[02-9]|[2-9][02-8]1|[2-9][02-8][02-9]))\s*(?:[.-]\s*)?)?([2-9]1[02-9]|[2-9][02-9]1|[2-9][02-9]{2})\s*(?:[.-]\s*)?([0-9]{4})(?:\s*(?:#|x\.?|ext\.?|extension)\s*(\d+))?'), text)
#    
#    if phone:
#        number = ''.join(phone[0])
#        if len(number) > 10:
#            return '+' + number
#        else:
#            return number
#
#
#def extract_mail(email):
#    email = re.findall("([^@|\s]+@[^@]+\.[^@|\s]+)", email)
#    #print(email)
#    if email:
#        try:
#            return email[0].split()[0].strip(';')
#        except IndexError:
#            return None
#
#def iter_headings(paragraphs):
#    for i in range(0,len(paragraphs)):
#        if paragraphs[i].style.name.startswith('Heading'):
#            yield paragraphs[i],i
#
#
#def extract_skills(doch,aft=[],skills=[]):
#    for heading,i in iter_headings(doch.paragraphs):
#        if 'Skills' in heading.text :
#            for j in range(i+1,len(doch.paragraphs)):
#                #print(doch.paragraphs[j].text)
#                aft.append(doch.paragraphs[j].text)
#                txt=doch.paragraphs[j].text
#                if doch.paragraphs[j].style.name.startswith('Heading'):
#                    break
#                if not txt:
#                    continue
#                else:
#                    if ':' in txt :
#                        word=txt.split(':')
#                        skills.append(word[1])
#                    else:
#                        skills.append(txt)    
#            return skills
#    

    
    
if __name__=="__main__":
    app.run(host='127.0.0.1', port=8080, debug=True)
    #execute(sys.argv)


























#Tried extracting skills using Bold letters but some of the words inside  paras also are in bold.
#Also tried using styles like heading ,normal ,list but some docs doesn'r represent with this format.
# tried usign font,caps ,small_caps but it doesn't have return text function
#tried with paras and tables extraction.
#now extracting name,phone number, email with regex and spacy.
#extracting skills with the help of styles and regex.


with open("C:\\Users\\P0142221\\ResParsing\\Data\\PranavKumar profile.docx",encoding="utf8",errors="ignore") as file:
    try:
        ert=file.readlines()
    except Exception as ex:
        print(ex)


# importing required modules 
import PyPDF2 

# creating a pdf file object 
pdfFileObj = open('C:\\Users\\P0142221\\ResParsing\\Data\\Chandra Prakash.pdf', 'rb') 

# creating a pdf reader object 
pdfReader = PyPDF2.PdfFileReader(pdfFileObj) 

# printing number of pages in pdf file 
print(pdfReader.numPages) 
rty=list()
for page in range(pdfReader.numPages):
    pageObj = pdfReader.getPage(page) 
    print(pageObj.values)
    rty.append(pageObj.extractText())
    
# creating a page object 


# closing the pdf file object 
pdfFileObj.close() 


from docx.api import Document
import docx
doc=Document("C:\\Users\\P0142221\\ResParsing\\Data\\Docx\\Akhil K.docx")

doch=Document("C:\\Users\\P0142221\\ResParsing\\Data\\Hyndavi java resume.docx")
doc1=docx.Document("C:\\Users\\P0142221\\ResParsing\\Data\\Sriram Data Science Profile.docx")
fullText = []
for para in doc.paragraphs:
    fullText.append(para.text)
full_t= ' '.join(fullText)
def extract_email(email):
    email = re.findall("([^@|\s]+@[^@]+\.[^@|\s]+)", email)
    print(email)
    if email:
        try:
            return email[0].split()[0].strip(';')
        except IndexError:
            return None
email = re.findall(re.compile("[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,4}"), full_t)
tab=doc.tables[0]
styl=doc.styles['Heading 1']
re=doc.styles
len(re)
docText = '\n\n'.join(
    paragraph.text for paragraph in doc.paragraphs
)

table=doc.tables[0]
table.rows
data = []

keys = None
for i, row in enumerate(table.rows):
    text = (cell.text for cell in row.cells)

    # Establish the mapping based on the first row
    # headers; these will become the keys of our dictionary
    if i == 0:
        keys = tuple(text)
        continue

    # Construct a dictionary for this row, mapping
    # keys to values for this row
    row_data = dict(zip(keys, text))
    data.append(row_data)


data[1]

paras=doc.paragraphs
len(paras)
para1=paras[7]


import pyodbc

server = 'tcp:idcdbserver.database.windows.net,1433'
database = 'IDCDB'
username = 'idclogin'
password = 'Idc@login'
#driver= '{ODBC Driver 17 for SQL Server}'
driver1='{SQL Server}'
key='DATA SCIENCE PROJECTS'
value='Prediction of Degree of Regulatory Compliance of an Insurer by analyzing the Complaints Resolution System(CRS). Classification of a transaction as fraudulent or not. Auto population of fields in a form by analysis of user selection patterns in the past.'
cnxn = pyodbc.connect('DRIVER='+driver1+';SERVER='+server+';DATABASE='+database+';UID='+username+';PWD='+ password)



try:
    cursor=cnxn.cursor()
    cursor.execute("""insert into doc_parsing(keys,value) 
                values(?,?)""",(key,value))
    print(cursor.commit())
    cursor.close()
    cnxn.close()
except Exception as e:
    print(e)
    

def azure_sql_connection():
    server = 'tcp:idcdbserver.database.windows.net,1433'
    database = 'IDCDB'
    username = 'idclogin'
    password = 'Idc@login'
    #driver= '{ODBC Driver 17 for SQL Server}'
    driver1='{SQL Server}'
    cnxn = pyodbc.connect('DRIVER='+driver1+';SERVER='+server+';DATABASE='+database+';UID='+username+';PWD='+ password)
    return cnxn

bolds=[]
italics=[]
for para in doc.paragraphs:
    for run in para.runs:
        print(run)
        if run.italic :
            italics.append(run.text)
        if run.bold :
            
            bolds.append(run.text)

boltalic_Dict={'bold_phrases':bolds,
              'italic_phrases':italics}

boltalic_Dict['bold_phrases'][0]

len(doc.tables)
sampl=False
ert=[]
bolds=[]
for para in doc.paragraphs:
    for run in para.runs:
        #print(run.text)
        if sampl:
            ert.append(run.text)
            
            #print(ert)
           
        if run.bold:
            #print(" ".join(ert))
            desc=" ".join(ert)
            if not bolds:
                key="Summary"
            else:
                
                key=bolds[-1]
            
                
                
            
            value=desc
            print(type(key))
            print(type(value))
            cnxn=azure_sql_connection()
            try:
                cursor=cnxn.cursor()
                cursor.execute("""insert into doc_parsing(keys,value) 
                values(?,?)""",(key,value))
                print(cursor.commit())
                cursor.close()
                cnxn.close()
            except Exception as e:
                print(e)
                
            sampl=True
            bolds.append(run.text)
            
            ert=[]
    
            
bolds[-1]

##*********************

doch=Document("C:\\Users\\P0142221\\ResParsing\\Data\\Tadasha_Nayak_Resume.docx")
len(doch.styles)
len(doch.paragraphs)
styl=doch.styles
Hd=styl['Heading 1']


for para in doch.paragraphs:
    for run in para.runs:
        print(run.text)

def iter_headings(paragraphs):
    for paragraph in paragraphs:
        if paragraph.style.name.startswith('Heading'):
            yield paragraph

for heading in iter_headings(doch.paragraphs):
    print(heading.text)


for paragraph in doch.paragraphs:
    for i in paragraph.runs:
        if run.small_caps:
            print(run.text)
            
###********************************
            
#bold
#style

doch=Document("C:\\Users\\P0142221\\ResParsing\\Data\\Sunjit Rana.G.docx")


def iter_headingss(paragraphs):
    for paragraph in paragraphs:
        if paragraph.style.name.startswith('Subtitle'):
            yield paragraph

def iter_headings(paragraphs):
    for i in range(0,len(paragraphs)):
        if paragraphs[i].style.name.startswith('Heading'):
            yield paragraphs[i],i

def iter_normal(paragraphs):
    for paragraph in paragraphs:
        if paragraph.style.name.startswith('Normal'):
            yield paragraph

def iter_list(paragraphs):
    for paragraph in paragraphs:
        if paragraph.style.name.startswith('List'):
            #print("Yes")
            yield paragraph

for heading in iter_normal(doch.paragraphs):
    print(heading.text)

from docx.enum.style import WD_STYLE_TYPE

styles = doch.styles
for s in styles:
    print(s.name)
    print(s.type)
paragraph_styles = [s for s in styles if s.type == WD_STYLE_TYPE.CHARACTER]
for style in paragraph_styles:
    print(style.name)
    
for para in doch.paragraphs:
    for run in para.runs:
        
aft=[]
skills=[]
for heading,i in iter_headings(doch.paragraphs):
    #print(heading.text)
    if 'Skills' in heading.text :
        for j in range(i+1,len(doch.paragraphs)):
            print(doch.paragraphs[j].text)
            aft.append(doch.paragraphs[j].text)
            if doch.paragraphs[j].style.name.startswith('Heading'):
                break
#            if not j:
#                continue
#            else:
#                if ':' in j :
#                    word=j.split(':')
#                    skills.append(word[1])
#                else:
#                    continue
            

for i in aft:
    if not i:
        print("dfa")
        continue
    else:
        if ':' in i:
            word=i.split(':')
            skills.append(word[1])
        else:
            continue
            
    



#Extracting skills with list
for heading in iter_headingss(doch.paragraphs):
    print(heading.text)
doch=Document("C:\\Users\\P0142221\\ResParsing\\Data\\Sriram Data Science Profile.docx")
  
#Extracting skills with tables

doch=Document("C:\\Users\\P0142221\\ResParsing\\Data\\Sriram Data Science Profile.docx")
len(doch.tables)
table=doch.tables[0]
table.rows
data = []

keys = None
for i, row in enumerate(table.rows):
    text = (cell.text for cell in row.cells)
    print(text)
    # Establish the mapping based on the first row
    # headers; these will become the keys of our dictionary
    if i == 0:
        keys = tuple(text)
        print(keys)
        continue

    # Construct a dictionary for this row, mapping
    # keys to values for this row
    row_data = dict(zip(keys, text))
    data.append(row_data)

        
for heading in iter_headingss(doch.paragraphs):
    #print(heading.text)
    if 'Skills' in heading.text :
        print(heading.text)        
          
import spacy
from spacy.matcher import Matcher 
       
nlp = spacy.load('en_core_web_sm')

matcher = Matcher(nlp.vocab)    
def extract_name(resume_text):
    nlp_text = nlp(resume_text)
    
    # First name and Last name are always Proper Nouns
    pattern = [{'POS': 'PROPN'}, {'POS': 'PROPN'}]
    
    matcher.add('NAME', None, *pattern)
    
    matches = matcher(nlp_text)
    
    for match_id, start, end in matches:
        span = nlp_text[start:end]
        return span.text

doch=Document("C:\\Users\\P0142221\\ResParsing\\Data\\Tadasha_Nayak_Resume.docx")
import docx2txt

temp = docx2txt.process("C:\\Users\\P0142221\\ResParsing\\Data\\Tadasha_Nayak_Resume.docx")
text = [line.replace('\t', ' ') for line in temp.split('\n') if line]
full_t= ' '.join(text)

text[0]

nlp_text = nlp(full_t)
pattern = [{'POS': 'PROPN'}, {'POS': 'PROPN'}]
    
matcher.add('NAME', None, pattern)
matches = matcher(nlp_text)
text[0:2]   
for match_id, start, end in matches:
    span = nlp_text[start:end]
    print(span.text)




import re

def extract_mobile_number(text):
    phone = re.findall(re.compile(r'(?:(?:\+?([1-9]|[0-9][0-9]|[0-9][0-9][0-9])\s*(?:[.-]\s*)?)?(?:\(\s*([2-9]1[02-9]|[2-9][02-8]1|[2-9][02-8][02-9])\s*\)|([0-9][1-9]|[0-9]1[02-9]|[2-9][02-8]1|[2-9][02-8][02-9]))\s*(?:[.-]\s*)?)?([2-9]1[02-9]|[2-9][02-9]1|[2-9][02-9]{2})\s*(?:[.-]\s*)?([0-9]{4})(?:\s*(?:#|x\.?|ext\.?|extension)\s*(\d+))?'), text)
    
    if phone:
        number = ''.join(phone[0])
        if len(number) > 10:
            return '+' + number
        else:
            return number
extract_mobile_number(full_t)

import re
def extract_email(email):
    email = re.findall("([^@|\s]+@[^@]+\.[^@|\s]+)", email)
    print(email)
    if email:
        try:
            return email[0].split()[0].strip(';')
        except IndexError:
            return None
        
extract_email(fullText)
email = re.findall("([^@|\s]+@[^@]+\.[^@|\s]+)", full_text)
txt=[]
for para in doch.paragraphs:
    txt.append(para.text)
full_tt= ' '.join(txt)
extract_email(full_tt)

### Skills extraction & Professional Summary Extraction & Number of Projects

## For Skills extraction we'll identify the styles , if there are styles then we
## we'll extract those styles and identify whether there is any technical skill 
## associated with that , if there are any technical skills then we'll extract
## the skills.
## If skills are not part of paras then we'll check the tables, if there are
## any tables , then we'll check the skills.
## If the above two things are not happening then we'll try and extract using
## list bullets and normal bullets.
## If still we are having issue then we;ll extract using spacy matcher with 
## the help of already pre defined skills.

#input from cmd = input()
#***********************
# Taking input from CMD

import sys
data=sys.argv[0]

# Conversion of doc to docx
import subprocess
filename="C:\\Users\\P0142221\\ResParsing\\Data\\karteek Data Scientist Profile.doc"
wee=subprocess.call(['lowriter', '--headless', '--convert-to', 'docx', filename])

# Conversion of pdf to docx

pdftodocx=subprocess.call('lowriter --invisible --convert-to doc "{}"'.format(filename), shell=True)

#The subprocess module enables you to start new applications from your Python program
#You can start a process in Python using the Popen function call. The program below starts the unix program ‘cat’ and the second parameter is the argument. This is equivalent to ‘cat test.py’
#process = Popen(['cat', 'test.py'], stdout=PIPE, stderr=PIPE)
#Subprocess has a method call() which can be used to start a program
#subprocess.call(["ls", "-l"])
#the full command would be “ls -l”
#code = subprocess.call(["ping", "www.yahoo.com"])
#code = subprocess.call("notepad.exe")

# Extracting Skills, Name , Phone Number , Email


#for top, dirs, files in os.walk('/my/pdf/folder'):
#    for filename in files:
#        if filename.endswith('.pdf'):
#            abspath = os.path.join(top, filename)
#            subprocess.call('lowriter --invisible --convert-to doc "{}"'
#                            .format(abspath), shell=True)

#  


@app.route('/upload',methods=['POST','GET'])
def parse_doc():
    if request.method=='POST':
        filename=request.files['img']
        #print(filename)
        txt=list()
        logging.getLogger().setLevel(logging.INFO)
        logging.info('File Name {}'.format(filename))
        fileextens=filename.filename.split('.')[1]
        print(fileextens)
        if fileextens=='docx':
            logging.info('File name ends with docx')
            doc=Document(filename)
            for para in doc.paragraphs:
                txt.append(para.text)
            full_t= ' '.join(txt)
            
            
            mob=ex.extract_mob_number(full_t)
            session['mob_session']=mob
            logging.info('Extracted Mobile  {}'.format(mob))
            
            mail=ex.extract_mail(full_t)
            session['mail_session']=mail
            logging.info('Extracted mail {}'.format(mail))
            print(session.get('mail_session'))
            skills=ex.extract_skills(doc)
            logging.info('Extracted Skills {}'.format(skills))
            session['skills_session']=skills
            
            
            if not mail or not name_e  or not mob  or not skills: 
                
                txt_doc=ex.txt_extraction(filename)
            if not mob:
                mob_doc=ex.extract_mob_number(txt_doc)
                print(mob_doc)
            if not mail:
                mail_doc=ex.extract_mail(txt_doc)
                print(mail_doc)
            if not name_e:
                name_doc=ex.name_extraction(txt_doc)
                print(name_doc)
            if not skills:
                
                url=ex.extract_linkedinurl(txt_doc)
                skills=linkedin.skills_linkdn(url)
                print(skills)
        print("skills")    
        return "Skills"
     
    return "File Not Uploaded"